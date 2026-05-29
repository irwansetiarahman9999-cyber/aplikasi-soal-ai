import streamlit as st
import PyPDF2
import google.generativeai as genai
import io
from docx import Document
from docx.shared import Pt

# ================= KONFIGURASI HALAMAN =================
st.set_page_config(page_title="Aplikasi Guru AI Pro - Bang Irwan", layout="wide")

# ================= KODE DESAIN BIRU PEKAT & KUNING PEKAT (CSS) =================
st.markdown("""
    <style>
    .stApp { background: linear-gradient(135deg, #051937 0%, #a27c0c 100%); }
    html, body, [data-testid="stHeader"], .stMarkdown, label { color: white !important; font-family: 'Arial'; }
    h1 { color: #ffc107 !important; text-shadow: 3px 3px 6px rgba(0,0,0,0.5); text-align: center; font-size: 40px !important; }
    h2, h3 { color: #ffc107 !important; border-bottom: 2px solid #ffc107; padding-bottom: 10px; }
    
    /* Warna area upload agar tidak samar */
    .stFileUploader { background-color: rgba(255, 255, 255, 0.15); border: 2px dashed #ffc107; border-radius: 15px; padding: 20px; }
    
    /* Desain Tombol Hijau Gakuk */
    .stButton>button {
        background-color: #1b5e20 !important;
        color: white !important;
        font-size: 18px !important;
        font-weight: bold !important;
        height: 60px !important;
        border-radius: 12px;
        border: 2px solid #ffc107;
        box-shadow: 0 4px 15px rgba(0,0,0,0.3);
        transition: 0.3s;
    }
    .stButton>button:hover { transform: scale(1.02); background-color: #2e7d32 !important; }
    
    /* Sidebar */
    [data-testid="stSidebar"] { background-color: #030e1d !important; }
    
    /* Kotak Identitas */
    .kop-wadah { background-color: rgba(0,0,0,0.2); padding: 20px; border-radius: 15px; margin-bottom: 20px; border: 1px solid rgba(255,255,255,0.1); }
    </style>
    """, unsafe_allow_html=True)

# ================= SIDEBAR =================
with st.sidebar:
    st.header("⚙️ Pengaturan AI")
    api_key = st.text_input("Gemini API Key:", type="password")
    st.divider()
    st.write("Dibuat khusus untuk Pak Guru Irwan Barambai.")

# ================= JUDUL UTAMA =================
st.title("APLIKASI ASISTEN GURU PRO (AI) 📄")
st.markdown("<p style='text-align: center; font-size: 20px;'>Solusi Cerdas Pembuatan Soal, Kisi-kisi, dan Rangkuman Materi</p>", unsafe_allow_html=True)
st.divider()

# ================= IDENTITAS & UPLOAD (SEJAJAR) =================
col_identitas, col_upload = st.columns([1.2, 1])

with col_identitas:
    st.header("1. Identitas Ujian")
    with st.container():
        st.markdown('<div class="kop-wadah">', unsafe_allow_html=True)
        sekolah = st.text_input("Nama Sekolah:", placeholder="Contoh: SDN Barambai Kolam Kiri 5")
        c_mapel, c_kelas = st.columns(2)
        mapel = c_mapel.text_input("Mata Pelajaran:")
        kelas = c_kelas.text_input("Kelas / Semester:")
        c_tgl, c_waktu = st.columns(2)
        tanggal = c_tgl.text_input("Tanggal Pelaksanaan:")
        waktu = c_waktu.text_input("Alokasi Waktu:")
        st.markdown('</div>', unsafe_allow_html=True)

with col_upload:
    st.header("2. Upload Materi (PDF)")
    file_pdf = st.file_uploader("Seret file PDF Anda ke sini:", type="pdf")
    if file_pdf:
        c_h1, c_h2 = st.columns(2)
        hal_awal = c_h1.number_input("Mulai Halaman:", min_value=1, value=1)
        hal_akhir = c_h2.number_input("Sampai Halaman:", min_value=1, value=5)

st.divider()

# ================= MATRIKS BLOOM (REVISI TERBARU) =================
st.header("3. Matriks Soal (Taksonomi Bloom)")
st.markdown("""
<div style="display: flex; justify-content: space-between; background-color: rgba(0,0,0,0.4); padding: 12px; border-radius: 8px; border: 1px solid #ffc107; margin-bottom: 10px;">
    <div style="flex: 2; font-weight: bold; color: #ffc107;">Tingkat Kognitif</div>
    <div style="flex: 1; text-align: center; font-weight: bold; color: #ffc107;">Pilihan Ganda</div>
    <div style="flex: 1; text-align: center; font-weight: bold; color: #ffc107;">Isian Singkat</div>
    <div style="flex: 1; text-align: center; font-weight: bold; color: #ffc107;">Uraian</div>
</div>
""", unsafe_allow_html=True)

tingkat_bloom = {
    "C1 - Mengingat (Remembering)": "C1",
    "C2 - Memahami (Understanding)": "C2",
    "C3 - Mengaplikasikan (Applying)": "C3",
    "C4 - Menganalisis (Analyzing)": "C4",
    "C5 - Evaluasi (Evaluating)": "C5",
    "C6 - Menciptakan (Creating)": "C6"
}

pg_data, isian_data, uraian_data = {}, {}, {}

for label, code in tingkat_bloom.items():
    c_label, c_pg, c_isi, c_ur = st.columns([2, 1, 1, 1])
    c_label.markdown(f"<div style='padding-top: 10px;'>{label}</div>", unsafe_allow_html=True)
    pg_data[code] = c_pg.number_input(f"PG {code}", 0, label_visibility="collapsed")
    isian_data[code] = c_isi.number_input(f"Isian {code}", 0, label_visibility="collapsed")
    uraian_data[code] = c_ur.number_input(f"Uraian {code}", 0, label_visibility="collapsed")

total_soal = sum(pg_data.values()) + sum(isian_data.values()) + sum(uraian_data.values())
st.info(f"Total Soal: **{total_soal}**")

st.divider()

# ================= 4 TOMBOL TERPISAH =================
st.header("4. Panel Eksekusi")
btn_col1, btn_col2, btn_col3, btn_col4 = st.columns(4)

tugas_ai = None
prompt_tambahan = ""

if btn_col1.button("📑 KISI-KISI SOAL"):
    tugas_ai = "KISI-KISI"
    prompt_tambahan = "Buatkan tabel Kisi-kisi soal yang mencakup: Materi, Indikator Soal, Level Kognitif, Bentuk Soal, dan Nomor Soal."

if btn_col2.button("📖 RANGKUMAN MATERI"):
    tugas_ai = "RANGKUMAN"
    prompt_tambahan = "Buatkan rangkuman materi yang esensial, padat, dan jelas berdasarkan teks PDF untuk dipelajari siswa."

if btn_col3.button("✍️ BUAT SOAL"):
    tugas_ai = "SOAL"
    prompt_tambahan = f"Buatkan soal ujian dengan rincian PG: {pg_data}, Isian: {isian_data}, Uraian: {uraian_data}. Wajib sertakan kode [C1-C6] di setiap soal dan Kunci Jawaban di akhir."

if btn_col4.button("🔥 CETAK SEMUA"):
    tugas_ai = "SEMUA"
    prompt_tambahan = f"Buatkan Rangkuman Materi, Kisi-kisi Soal, dan Soal Ujian (PG: {pg_data}, Isian: {isian_data}, Uraian: {uraian_data}) lengkap dengan Kunci Jawaban."

# ================= PROSES AI =================
if tugas_ai:
    if not api_key or not file_pdf:
        st.error("⚠️ Silakan masukkan API Key dan unggah file PDF dulu, Bang!")
    else:
        with st.spinner(f'Sedang memproses {tugas_ai}... Mohon tunggu sebentar.'):
            try:
                # Konfigurasi AI & Pelacak Model
                genai.configure(api_key=api_key)
                model_name = 'gemini-1.5-flash'
                for m in genai.list_models():
                    if 'generateContent' in m.supported_generation_methods:
                        if 'flash' in m.name or 'pro' in m.name:
                            model_name = m.name.replace('models/', '')
                            break
                model = genai.GenerativeModel(model_name)

                # Ekstraksi PDF
                pdf_reader = PyPDF2.PdfReader(file_pdf)
                teks_pdf = ""
                for i in range(max(0, hal_awal-1), min(len(pdf_reader.pages), hal_akhir)):
                    teks_pdf += pdf_reader.pages[i].extract_text() + "\n"

                # Jalankan Perintah
                kop_soal = f"Sekolah: {sekolah}\nMapel: {mapel}\nKelas: {kelas}\nTanggal: {tanggal}\nWaktu: {waktu}\n"
                full_prompt = f"Anda adalah Asisten Guru Profesional. {prompt_tambahan}\n\nKOP SOAL:\n{kop_soal}\n\nMATERI:\n{teks_pdf}"
                
                response = model.generate_content(full_prompt)
                
                # Tampilkan Hasil
                st.success(f"Berhasil! Hasil {tugas_ai} telah siap.")
                st.markdown(f'<div style="background-color:rgba(0,0,0,0.4); padding:25px; border-radius:15px; border:1px solid #ffc107;">{response.text}</div>', unsafe_allow_html=True)
                
                # Tombol Download Word
                doc = Document()
                doc.add_heading(f"HASIL {tugas_ai} - {mapel}", 0)
                doc.add_paragraph(kop_soal)
                doc.add_paragraph(response.text)
                
                buf = io.BytesIO()
                doc.save(buf)
                st.download_button(
                    label=f"📥 DOWNLOAD HASIL {tugas_ai} (.docx)",
                    data=buf.getvalue(),
                    file_name=f"{tugas_ai}_{mapel}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Terjadi kesalahan: {e}")